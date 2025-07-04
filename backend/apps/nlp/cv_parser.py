"""
Advanced CV Parser - huntRED® v2
Parser súper potente de CVs con NLP, ML y múltiples formatos.
"""

import re
import json
import spacy
import fitz  # PyMuPDF
import docx
import pandas as pd
from typing import Dict, List, Optional, Any, Tuple
from datetime import datetime, timedelta
from dataclasses import dataclass, asdict
import numpy as np
from transformers import pipeline, AutoTokenizer, AutoModel
import torch
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import logging
import email
import io
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import phonenumbers
from dateutil import parser as date_parser
import requests
import base64

logger = logging.getLogger(__name__)


@dataclass
class ParsedCV:
    """Estructura de datos para CV parseado."""
    # Información personal
    full_name: str
    email: Optional[str]
    phone: Optional[str]
    address: Optional[str]
    linkedin_url: Optional[str]
    github_url: Optional[str]
    portfolio_url: Optional[str]
    
    # Información profesional
    current_title: Optional[str]
    summary: Optional[str]
    years_of_experience: Optional[int]
    
    # Experiencia laboral
    experience: List[Dict]
    
    # Educación
    education: List[Dict]
    
    # Habilidades
    technical_skills: List[str]
    soft_skills: List[str]
    languages: List[Dict]
    
    # Certificaciones y logros
    certifications: List[Dict]
    awards: List[Dict]
    publications: List[Dict]
    
    # Información adicional
    keywords: List[str]
    industry_match: List[str]
    confidence_score: float
    processing_time: float
    
    # Metadata
    source_file: Optional[str]
    file_type: Optional[str]
    parsed_at: datetime
    parser_version: str


class SuperCVParser:
    """
    Parser súper potente de CVs con múltiples tecnologías de NLP y ML.
    """
    
    def __init__(self, config: Dict[str, Any]):
        self.config = config
        self.logger = logger
        
        # Cargar modelos NLP
        self._load_nlp_models()
        
        # Configurar extractores
        self._setup_extractors()
        
        # Configurar clasificadores
        self._setup_classifiers()
        
        # Patrones de expresiones regulares
        self._setup_regex_patterns()
        
    def _load_nlp_models(self):
        """Carga todos los modelos de NLP necesarios."""
        try:
            # spaCy models
            self.nlp_en = spacy.load("en_core_web_lg")
            self.nlp_es = spacy.load("es_core_news_lg")
            
            # BERT para análisis semántico
            self.bert_tokenizer = AutoTokenizer.from_pretrained('bert-base-uncased')
            self.bert_model = AutoModel.from_pretrained('bert-base-uncased')
            
            # Pipeline para clasificación de texto
            self.classifier = pipeline(
                "text-classification",
                model="microsoft/DialoGPT-medium",
                tokenizer="microsoft/DialoGPT-medium"
            )
            
            # Pipeline para extracción de entidades
            self.ner_pipeline = pipeline(
                "ner",
                model="dbmdz/bert-large-cased-finetuned-conll03-english",
                aggregation_strategy="simple"
            )
            
            # TF-IDF Vectorizer para análisis de skills
            self.tfidf_vectorizer = TfidfVectorizer(
                max_features=5000,
                stop_words='english',
                ngram_range=(1, 3)
            )
            
            logger.info("NLP models loaded successfully")
            
        except Exception as e:
            logger.error(f"Error loading NLP models: {str(e)}")
            raise
    
    def _setup_extractors(self):
        """Configura extractores especializados."""
        # Base de datos de skills técnicas
        self.technical_skills_db = {
            'programming_languages': [
                'python', 'java', 'javascript', 'c++', 'c#', 'ruby', 'php',
                'go', 'rust', 'kotlin', 'swift', 'typescript', 'scala',
                'r', 'matlab', 'perl', 'shell', 'bash', 'powershell'
            ],
            'frameworks': [
                'django', 'flask', 'fastapi', 'react', 'angular', 'vue',
                'node.js', 'express', 'spring', 'laravel', 'rails',
                'tensorflow', 'pytorch', 'keras', 'scikit-learn'
            ],
            'databases': [
                'mysql', 'postgresql', 'mongodb', 'redis', 'elasticsearch',
                'oracle', 'sql server', 'sqlite', 'cassandra', 'neo4j'
            ],
            'cloud_platforms': [
                'aws', 'azure', 'gcp', 'google cloud', 'digitalocean',
                'heroku', 'vercel', 'netlify', 'cloudflare'
            ],
            'tools': [
                'docker', 'kubernetes', 'jenkins', 'git', 'github',
                'gitlab', 'jira', 'confluence', 'slack', 'trello'
            ]
        }
        
        # Soft skills database
        self.soft_skills_db = [
            'leadership', 'communication', 'teamwork', 'problem solving',
            'analytical thinking', 'creativity', 'adaptability', 'time management',
            'project management', 'negotiation', 'presentation', 'mentoring'
        ]
        
        # Patrones de industrias
        self.industry_patterns = {
            'technology': ['software', 'tech', 'it', 'computer', 'digital'],
            'finance': ['bank', 'finance', 'accounting', 'investment'],
            'healthcare': ['health', 'medical', 'hospital', 'clinic'],
            'education': ['education', 'teaching', 'university', 'school'],
            'marketing': ['marketing', 'advertising', 'brand', 'social media']
        }
    
    def _setup_classifiers(self):
        """Configura clasificadores ML."""
        # Clasificador de nivel de experiencia
        self.experience_levels = {
            'entry': {'min_years': 0, 'max_years': 2, 'keywords': ['junior', 'entry', 'trainee']},
            'mid': {'min_years': 2, 'max_years': 5, 'keywords': ['mid', 'intermediate', 'developer']},
            'senior': {'min_years': 5, 'max_years': 10, 'keywords': ['senior', 'lead', 'principal']},
            'expert': {'min_years': 10, 'max_years': None, 'keywords': ['expert', 'architect', 'director']}
        }
    
    def _setup_regex_patterns(self):
        """Configura patrones de expresiones regulares."""
        self.patterns = {
            'email': r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
            'phone': r'(\+?\d{1,3}[-.\s]?)?\(?[\d\s-]{7,15}\)?',
            'linkedin': r'linkedin\.com\/in\/[\w\-\.]+',
            'github': r'github\.com\/[\w\-\.]+',
            'url': r'https?://(?:[-\w.])+(?:[:\d]+)?(?:/(?:[\w/_.])*(?:\?(?:[\w&=%.])*)?(?:#(?:[\w.])*)?)?',
            'years_experience': r'(\d+)[\+\s]*(?:years?|yrs?)\s*(?:of\s*)?(?:experience|exp)',
            'date_ranges': r'(\d{4})\s*[-–]\s*(\d{4}|present|current)',
            'gpa': r'(?:gpa|grade point average)[:]\s*(\d+\.?\d*)',
            'degree': r'(bachelor|master|phd|doctorate|mba|b\.?\s*[aes]\.?|m\.?\s*[aes]\.?|ph\.?\s*d\.?)',
        }
        
        # Compilar patrones
        self.compiled_patterns = {
            name: re.compile(pattern, re.IGNORECASE)
            for name, pattern in self.patterns.items()
        }
    
    def parse_cv_from_file(self, file_path: str, file_type: str = None) -> ParsedCV:
        """Parser principal para archivos de CV."""
        start_time = datetime.now()
        
        try:
            # Detectar tipo de archivo si no se proporciona
            if not file_type:
                file_type = self._detect_file_type(file_path)
            
            # Extraer texto según el tipo de archivo
            raw_text = self._extract_text_from_file(file_path, file_type)
            
            # Procesar el texto extraído
            parsed_cv = self._process_cv_text(raw_text)
            
            # Añadir metadata
            parsed_cv.source_file = file_path
            parsed_cv.file_type = file_type
            parsed_cv.processing_time = (datetime.now() - start_time).total_seconds()
            
            return parsed_cv
            
        except Exception as e:
            logger.error(f"Error parsing CV file {file_path}: {str(e)}")
            raise
    
    def parse_cv_from_email(self, email_content: str) -> ParsedCV:
        """Parser para CVs recibidos por email."""
        start_time = datetime.now()
        
        try:
            # Parsear email
            msg = email.message_from_string(email_content)
            
            # Extraer texto del cuerpo del email
            email_text = self._extract_text_from_email(msg)
            
            # Buscar adjuntos con CVs
            cv_text = email_text
            for part in msg.walk():
                if part.get_content_disposition() == 'attachment':
                    filename = part.get_filename()
                    if filename and self._is_cv_file(filename):
                        # Procesar adjunto
                        attachment_data = part.get_payload(decode=True)
                        attachment_text = self._extract_text_from_bytes(
                            attachment_data, 
                            self._detect_file_type(filename)
                        )
                        cv_text += "\n\n" + attachment_text
            
            # Procesar el texto combinado
            parsed_cv = self._process_cv_text(cv_text)
            parsed_cv.source_file = "email"
            parsed_cv.file_type = "email"
            parsed_cv.processing_time = (datetime.now() - start_time).total_seconds()
            
            return parsed_cv
            
        except Exception as e:
            logger.error(f"Error parsing CV from email: {str(e)}")
            raise
    
    def parse_cv_from_chatbot(self, chat_messages: List[Dict]) -> ParsedCV:
        """Parser para CVs enviados por chatbot."""
        start_time = datetime.now()
        
        try:
            # Combinar todos los mensajes del chat
            combined_text = ""
            for message in chat_messages:
                if message.get('type') == 'text':
                    combined_text += message.get('content', '') + "\n"
                elif message.get('type') == 'file':
                    # Procesar archivo enviado por chat
                    file_data = message.get('file_data')
                    file_type = message.get('file_type')
                    if file_data and self._is_cv_file(message.get('filename', '')):
                        file_text = self._extract_text_from_bytes(
                            base64.b64decode(file_data), 
                            file_type
                        )
                        combined_text += "\n\n" + file_text
            
            # Procesar el texto
            parsed_cv = self._process_cv_text(combined_text)
            parsed_cv.source_file = "chatbot"
            parsed_cv.file_type = "chat"
            parsed_cv.processing_time = (datetime.now() - start_time).total_seconds()
            
            return parsed_cv
            
        except Exception as e:
            logger.error(f"Error parsing CV from chatbot: {str(e)}")
            raise
    
    def _process_cv_text(self, text: str) -> ParsedCV:
        """Procesa el texto del CV y extrae toda la información."""
        try:
            # Limpiar y normalizar el texto
            cleaned_text = self._clean_text(text)
            
            # Detectar idioma
            language = self._detect_language(cleaned_text)
            nlp = self.nlp_es if language == 'es' else self.nlp_en
            
            # Procesar con spaCy
            doc = nlp(cleaned_text)
            
            # Extraer información básica
            personal_info = self._extract_personal_info(cleaned_text, doc)
            contact_info = self._extract_contact_info(cleaned_text)
            
            # Extraer experiencia laboral
            experience = self._extract_work_experience(cleaned_text, doc)
            
            # Extraer educación
            education = self._extract_education(cleaned_text, doc)
            
            # Extraer habilidades
            technical_skills, soft_skills = self._extract_skills(cleaned_text, doc)
            
            # Extraer idiomas
            languages = self._extract_languages(cleaned_text, doc)
            
            # Extraer certificaciones
            certifications = self._extract_certifications(cleaned_text, doc)
            
            # Extraer awards y publicaciones
            awards = self._extract_awards(cleaned_text, doc)
            publications = self._extract_publications(cleaned_text, doc)
            
            # Análisis semántico avanzado
            keywords = self._extract_keywords(cleaned_text)
            industry_match = self._classify_industry(cleaned_text)
            
            # Calcular años de experiencia
            years_experience = self._calculate_years_of_experience(experience, cleaned_text)
            
            # Calcular confidence score
            confidence_score = self._calculate_confidence_score(
                personal_info, contact_info, experience, education, technical_skills
            )
            
            # Crear objeto ParsedCV
            parsed_cv = ParsedCV(
                # Información personal
                full_name=personal_info.get('name', ''),
                email=contact_info.get('email'),
                phone=contact_info.get('phone'),
                address=contact_info.get('address'),
                linkedin_url=contact_info.get('linkedin'),
                github_url=contact_info.get('github'),
                portfolio_url=contact_info.get('portfolio'),
                
                # Información profesional
                current_title=personal_info.get('current_title'),
                summary=personal_info.get('summary'),
                years_of_experience=years_experience,
                
                # Experiencia y educación
                experience=experience,
                education=education,
                
                # Habilidades
                technical_skills=technical_skills,
                soft_skills=soft_skills,
                languages=languages,
                
                # Certificaciones y logros
                certifications=certifications,
                awards=awards,
                publications=publications,
                
                # Análisis
                keywords=keywords,
                industry_match=industry_match,
                confidence_score=confidence_score,
                processing_time=0.0,  # Se actualizará después
                
                # Metadata
                source_file=None,
                file_type=None,
                parsed_at=datetime.now(),
                parser_version="2.0.0"
            )
            
            return parsed_cv
            
        except Exception as e:
            logger.error(f"Error processing CV text: {str(e)}")
            raise
    
    def _extract_text_from_file(self, file_path: str, file_type: str) -> str:
        """Extrae texto de diferentes tipos de archivo."""
        try:
            if file_type.lower() == 'pdf':
                return self._extract_from_pdf(file_path)
            elif file_type.lower() in ['doc', 'docx']:
                return self._extract_from_docx(file_path)
            elif file_type.lower() == 'txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extrae texto de PDF usando PyMuPDF."""
        try:
            doc = fitz.open(file_path)
            text = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text += page.get_text()
            
            doc.close()
            return text
            
        except Exception as e:
            logger.error(f"Error extracting PDF text: {str(e)}")
            raise
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extrae texto de DOCX usando python-docx."""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extraer texto de tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {str(e)}")
            raise
    
    def _extract_text_from_bytes(self, data: bytes, file_type: str) -> str:
        """Extrae texto de datos binarios."""
        try:
            if file_type.lower() == 'pdf':
                doc = fitz.open(stream=data, filetype="pdf")
                text = ""
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    text += page.get_text()
                doc.close()
                return text
            elif file_type.lower() in ['doc', 'docx']:
                doc = docx.Document(io.BytesIO(data))
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                return text
            else:
                return data.decode('utf-8', errors='ignore')
                
        except Exception as e:
            logger.error(f"Error extracting text from bytes: {str(e)}")
            return ""
    
    def _extract_personal_info(self, text: str, doc) -> Dict[str, Any]:
        """Extrae información personal básica."""
        info = {}
        
        # Extraer nombre usando NER
        persons = [ent.text for ent in doc.ents if ent.label_ == "PERSON"]
        if persons:
            # Tomar el primer nombre que aparece
            info['name'] = persons[0]
        else:
            # Fallback: buscar en las primeras líneas
            lines = text.split('\n')[:5]
            for line in lines:
                if re.match(r'^[A-Z][a-z]+ [A-Z][a-z]+', line.strip()):
                    info['name'] = line.strip()
                    break
        
        # Extraer título profesional actual
        title_patterns = [
            r'(?:^|\n)([A-Z][a-z]+ [A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s*(?:\n|$)',
            r'(?:title|position|role)[:]\s*(.+?)(?:\n|$)',
            r'(?:currently|current).*?(?:working as|position)[:]\s*(.+?)(?:\n|$)'
        ]
        
        for pattern in title_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                info['current_title'] = match.group(1).strip()
                break
        
        # Extraer resumen/objetivo
        summary_patterns = [
            r'(?:summary|objective|profile|about)[:]\s*(.+?)(?:\n\n|\n[A-Z])',
            r'(?:professional summary)[:]\s*(.+?)(?:\n\n|\n[A-Z])'
        ]
        
        for pattern in summary_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                info['summary'] = match.group(1).strip()
                break
        
        return info
    
    def _extract_contact_info(self, text: str) -> Dict[str, Optional[str]]:
        """Extrae información de contacto."""
        contact = {}
        
        # Email
        email_match = self.compiled_patterns['email'].search(text)
        contact['email'] = email_match.group(0) if email_match else None
        
        # Teléfono
        phone_matches = self.compiled_patterns['phone'].findall(text)
        if phone_matches:
            # Validar y formatear teléfono
            for phone in phone_matches:
                try:
                    parsed_phone = phonenumbers.parse(phone, None)
                    if phonenumbers.is_valid_number(parsed_phone):
                        contact['phone'] = phonenumbers.format_number(
                            parsed_phone, 
                            phonenumbers.PhoneNumberFormat.INTERNATIONAL
                        )
                        break
                except:
                    continue
        
        # LinkedIn
        linkedin_match = self.compiled_patterns['linkedin'].search(text)
        if linkedin_match:
            contact['linkedin'] = f"https://{linkedin_match.group(0)}"
        
        # GitHub
        github_match = self.compiled_patterns['github'].search(text)
        if github_match:
            contact['github'] = f"https://{github_match.group(0)}"
        
        # Portfolio/Website
        url_matches = self.compiled_patterns['url'].findall(text)
        for url in url_matches:
            if 'linkedin' not in url.lower() and 'github' not in url.lower():
                contact['portfolio'] = url
                break
        
        # Dirección (usando NER para GPE - Geopolitical entities)
        # Simplificado para este ejemplo
        address_patterns = [
            r'(?:address|location)[:]\s*(.+?)(?:\n|$)',
            r'(\d+\s+[A-Za-z\s]+(?:street|st|avenue|ave|road|rd|drive|dr|boulevard|blvd).*?)(?:\n|$)'
        ]
        
        for pattern in address_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                contact['address'] = match.group(1).strip()
                break
        
        return contact
    
    def _extract_work_experience(self, text: str, doc) -> List[Dict]:
        """Extrae experiencia laboral."""
        experiences = []
        
        # Buscar sección de experiencia
        experience_section = self._extract_section(text, 'experience|work|employment|career')
        
        if not experience_section:
            return experiences
        
        # Dividir en experiencias individuales
        # Buscar patrones que indiquen nuevas experiencias
        job_patterns = [
            r'(\d{4})\s*[-–]\s*(\d{4}|present|current)\s*[:]\s*(.+?)(?=\d{4}|$)',
            r'([A-Z][^,\n]+),\s*([A-Z][^,\n]+)\s*\((\d{4})\s*[-–]\s*(\d{4}|present)\)',
        ]
        
        # NER para organizaciones
        orgs = [ent.text for ent in doc.ents if ent.label_ == "ORG"]
        
        # Método más sofisticado usando estructura del texto
        lines = experience_section.split('\n')
        current_job = {}
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Detectar título de trabajo
            if re.match(r'^[A-Z][a-z]+(?: [A-Z][a-z]+)*$', line):
                if current_job:
                    experiences.append(current_job.copy())
                current_job = {'title': line}
            
            # Detectar empresa
            elif any(org.lower() in line.lower() for org in orgs):
                current_job['company'] = line
            
            # Detectar fechas
            elif self.compiled_patterns['date_ranges'].search(line):
                date_match = self.compiled_patterns['date_ranges'].search(line)
                if date_match:
                    current_job['start_date'] = date_match.group(1)
                    current_job['end_date'] = date_match.group(2)
            
            # Descripción (líneas que no coinciden con patrones anteriores)
            else:
                if 'description' not in current_job:
                    current_job['description'] = line
                else:
                    current_job['description'] += ' ' + line
        
        # Añadir último trabajo
        if current_job:
            experiences.append(current_job)
        
        return experiences
    
    def _extract_education(self, text: str, doc) -> List[Dict]:
        """Extrae información educativa."""
        education = []
        
        # Buscar sección de educación
        education_section = self._extract_section(text, 'education|academic|qualifications|studies')
        
        if not education_section:
            return education
        
        # Buscar títulos y universidades
        degree_matches = self.compiled_patterns['degree'].finditer(education_section)
        orgs = [ent.text for ent in doc.ents if ent.label_ == "ORG"]
        
        lines = education_section.split('\n')
        current_edu = {}
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Detectar grado
            if self.compiled_patterns['degree'].search(line):
                if current_edu:
                    education.append(current_edu.copy())
                current_edu = {'degree': line}
            
            # Detectar institución
            elif any(org.lower() in line.lower() for org in orgs):
                current_edu['institution'] = line
            
            # Detectar fechas
            elif self.compiled_patterns['date_ranges'].search(line):
                date_match = self.compiled_patterns['date_ranges'].search(line)
                if date_match:
                    current_edu['start_date'] = date_match.group(1)
                    current_edu['end_date'] = date_match.group(2)
            
            # Detectar GPA
            elif self.compiled_patterns['gpa'].search(line):
                gpa_match = self.compiled_patterns['gpa'].search(line)
                current_edu['gpa'] = float(gpa_match.group(1))
        
        if current_edu:
            education.append(current_edu)
        
        return education
    
    def _extract_skills(self, text: str, doc) -> Tuple[List[str], List[str]]:
        """Extrae habilidades técnicas y blandas."""
        text_lower = text.lower()
        
        technical_skills = []
        soft_skills = []
        
        # Extraer habilidades técnicas
        for category, skills in self.technical_skills_db.items():
            for skill in skills:
                if skill.lower() in text_lower:
                    technical_skills.append(skill)
        
        # Extraer soft skills
        for skill in self.soft_skills_db:
            if skill.lower() in text_lower:
                soft_skills.append(skill)
        
        # Análisis semántico adicional con BERT
        technical_skills.extend(self._extract_skills_with_bert(text, 'technical'))
        soft_skills.extend(self._extract_skills_with_bert(text, 'soft'))
        
        # Remover duplicados
        technical_skills = list(set(technical_skills))
        soft_skills = list(set(soft_skills))
        
        return technical_skills, soft_skills
    
    def _extract_skills_with_bert(self, text: str, skill_type: str) -> List[str]:
        """Extrae habilidades usando análisis semántico con BERT."""
        try:
            # Dividir texto en chunks para BERT
            max_length = 512
            chunks = [text[i:i+max_length] for i in range(0, len(text), max_length)]
            
            skills = []
            for chunk in chunks[:3]:  # Limitar a 3 chunks para performance
                # Tokenizar y procesar con BERT
                inputs = self.bert_tokenizer(chunk, return_tensors="pt", truncation=True, padding=True)
                
                with torch.no_grad():
                    outputs = self.bert_model(**inputs)
                    embeddings = outputs.last_hidden_state.mean(dim=1)
                
                # Aquí implementarías la lógica de clasificación específica
                # Por simplicidad, retornamos lista vacía
                pass
            
            return skills
            
        except Exception as e:
            logger.error(f"Error in BERT skill extraction: {str(e)}")
            return []
    
    def _extract_section(self, text: str, section_pattern: str) -> str:
        """Extrae una sección específica del CV."""
        pattern = rf'(?:^|\n)({section_pattern})[:]*\s*\n(.*?)(?=\n\n[A-Z]|\n[A-Z][A-Z]|$)'
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL | re.MULTILINE)
        
        if match:
            return match.group(2).strip()
        
        return ""
    
    def _detect_language(self, text: str) -> str:
        """Detecta el idioma del texto."""
        # Patrones simples para detectar español vs inglés
        spanish_words = ['experiencia', 'educación', 'habilidades', 'trabajo', 'empresa']
        english_words = ['experience', 'education', 'skills', 'work', 'company']
        
        spanish_count = sum(1 for word in spanish_words if word in text.lower())
        english_count = sum(1 for word in english_words if word in text.lower())
        
        return 'es' if spanish_count > english_count else 'en'
    
    def _clean_text(self, text: str) -> str:
        """Limpia y normaliza el texto."""
        # Remover caracteres especiales excesivos
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[^\w\s\-\.\,\(\)\@\:\;]', ' ', text)
        
        # Normalizar espacios
        text = '\n'.join(line.strip() for line in text.split('\n') if line.strip())
        
        return text
    
    def _calculate_confidence_score(self, personal_info: Dict, contact_info: Dict, 
                                  experience: List, education: List, skills: List) -> float:
        """Calcula un score de confianza basado en la completitud de los datos."""
        score = 0.0
        max_score = 10.0
        
        # Información personal (20%)
        if personal_info.get('name'):
            score += 2.0
        
        # Información de contacto (20%)
        if contact_info.get('email'):
            score += 1.0
        if contact_info.get('phone'):
            score += 1.0
        
        # Experiencia (30%)
        if experience:
            score += 2.0
            if len(experience) >= 3:
                score += 1.0
        
        # Educación (20%)
        if education:
            score += 1.5
            if len(education) >= 2:
                score += 0.5
        
        # Habilidades (10%)
        if skills and len(skills) >= 5:
            score += 1.0
        
        return min(score / max_score, 1.0)
    
    # Métodos adicionales para completar el parser...
    def _extract_languages(self, text: str, doc) -> List[Dict]:
        """Extrae idiomas."""
        return []  # Implementación simplificada
    
    def _extract_certifications(self, text: str, doc) -> List[Dict]:
        """Extrae certificaciones."""
        return []  # Implementación simplificada
    
    def _extract_awards(self, text: str, doc) -> List[Dict]:
        """Extrae premios y reconocimientos."""
        return []  # Implementación simplificada
    
    def _extract_publications(self, text: str, doc) -> List[Dict]:
        """Extrae publicaciones."""
        return []  # Implementación simplificada
    
    def _extract_keywords(self, text: str) -> List[str]:
        """Extrae keywords relevantes."""
        return []  # Implementación simplificada
    
    def _classify_industry(self, text: str) -> List[str]:
        """Clasifica la industria del candidato."""
        return []  # Implementación simplificada
    
    def _calculate_years_of_experience(self, experience: List, text: str) -> Optional[int]:
        """Calcula años totales de experiencia."""
        return None  # Implementación simplificada
    
    def _detect_file_type(self, file_path: str) -> str:
        """Detecta el tipo de archivo."""
        extension = file_path.split('.')[-1].lower()
        return extension
    
    def _is_cv_file(self, filename: str) -> bool:
        """Verifica si un archivo es un CV."""
        cv_extensions = ['pdf', 'doc', 'docx', 'txt']
        extension = filename.split('.')[-1].lower()
        return extension in cv_extensions
    
    def _extract_text_from_email(self, msg) -> str:
        """Extrae texto del cuerpo de un email."""
        if msg.is_multipart():
            for part in msg.walk():
                if part.get_content_type() == "text/plain":
                    return part.get_payload(decode=True).decode()
        else:
            return msg.get_payload(decode=True).decode()
        return ""