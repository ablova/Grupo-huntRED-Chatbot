# /home/pablo/app/ats/utils/parser.py
from typing import Dict, List, Optional, Any, Tuple, Union
import sys
import logging
import unicodedata
import email
import asyncio
import re
import concurrent.futures
import io
from concurrent.futures import ThreadPoolExecutor, as_completed
from functools import lru_cache
import json
from pathlib import Path
from datetime import datetime
from urllib.parse import urljoin, urlparse
from tempfile import NamedTemporaryFile
import unicodedata
from bs4 import BeautifulSoup
import trafilatura
from app.ats.utils.scraping.scraping_utils import PlaywrightAntiDeteccion
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
from spacy.matcher import Matcher
from spacy.tokens import Doc
import psutil
import gc
# from app.tasks import process_message  # Import roto, comentado temporalmente
from django.conf import settings
from django.core.exceptions import ValidationError
from dataclasses import dataclass
# from app.ats.utils.text_processing import TextProcessor  # Import roto, comentado temporalmente
# from app.ats.utils.validation import DocumentValidator  # Import roto, comentado temporalmente

# Configure logger
logger = logging.getLogger(__name__)
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler('app.log')  # Log to file for debugging
    ]
)

# Check for resume-parser
try:
    import resume_parser as resumeparse
    HAS_RESUME_PARSER = True
    logger.info("resume_parser available, using en_core_web_sm")
except ImportError:
    HAS_RESUME_PARSER = False
    logger.warning("resume-parser not installed. Some features may be limited.")

# External libraries
try:
    import trafilatura
    import pdfplumber
    from docx import Document
    from langdetect import detect, LangDetectException
    from langdetect.lang_detect_exception import LangDetectException as LDE
    import spacy
    from spacy.language import Language
    from spacy_langdetect import LanguageDetector
    from aioimaplib import aioimaplib
    from contextlib import ExitStack
    from django.utils.timezone import now
    from django.db import transaction
    from django.core.cache import cache
    from django.core.mail import send_mail
    from django.template.loader import render_to_string
    from asgiref.sync import sync_to_async
    from app.ats.chatbot.utils.nlp_utils import NLPUtils
    from app.ats.chatbot.nlp.nlp import NLPProcessor
    from app.models import ConfiguracionBU, Person, BusinessUnit, Division, Skill, Conversation, Vacante
    from app.ats.integrations.services import EmailService, MessageService
    from app.ats.chatbot.components.chat_state_manager import ChatStateManager
    from app.ats.chatbot.workflow.common.common import get_possible_transitions, process_business_unit_transition
    from app.ats.chatbot.workflow.business_units.huntred.huntred import process_huntred_candidate
    from app.ats.chatbot.workflow.business_units.huntu.huntu import process_huntu_candidate
    from app.ats.chatbot.workflow.business_units.amigro.amigro import process_amigro_candidate
    from app.ats.chatbot.workflow.business_units.sexsi.sexsi import process_sexsi_payment
    from app.ats.utils.report_generator import ReportGenerator
except ImportError as e:
    logger.error(f"Missing required package: {e}")
    raise

# Configuration
DEFAULT_LANGUAGE = 'es'
MAX_WORKERS = 4
CACHE_TIMEOUT = 60 * 60 * 24  # 24 hours
JOB_KEYWORDS = ['job', 'vacancy', 'position', 'opening', 'empleo', 'vacante', 'puesto']

# Global cache for division skills
DIVISION_SKILLS_CACHE = None

# SpaCy language models
NLP_MODELS = {}

def get_nlp(lang: str = 'es') -> Language:
    """Get or load a spaCy language model with caching."""
    if lang not in NLP_MODELS:
        try:
            if lang == 'es':
                NLP_MODELS[lang] = spacy.load('es_core_news_lg')
            else:
                NLP_MODELS[lang] = spacy.load('en_core_web_lg')
            if 'language_detector' not in NLP_MODELS[lang].pipe_names:
                def create_lang_detector(nlp, name):
                    return LanguageDetector()
                Language.factory('language_detector', func=create_lang_detector)
                NLP_MODELS[lang].add_pipe('language_detector', last=True)
        except OSError:
            logger.warning(f"SpaCy model for {lang} not found, downloading...")
            try:
                import subprocess
                subprocess.run([sys.executable, "-m", "spacy", "download",
                              f"{lang}_core_news_lg" if lang == 'es' else "en_core_web_lg"],
                             check=True)
                return get_nlp(lang)
            except Exception as e:
                logger.error(f"Failed to download spaCy model: {e}")
                return None
    return NLP_MODELS[lang]

# IMAP folder configuration
FOLDER_CONFIG = {
    "inbox": "INBOX",
    "cv_folder": "INBOX.CV",
    "parsed_folder": "INBOX.Parsed",
    "error_folder": "INBOX.Error",
}

@lru_cache(maxsize=1024)
def detect_language(text: str, default: str = DEFAULT_LANGUAGE) -> str:
    """Detect the language of the given text with caching."""
    if not text or not text.strip():
        return default
    cache_key = f'lang_detect:{hash(text[:200])}'
    cached = cache.get(cache_key)
    if cached:
        return cached
    try:
        nlp = get_nlp('es')
        if nlp:
            doc = nlp(text[:5000])
            if hasattr(doc, '_.language'):
                lang = doc._.language.get('language')
                if lang in ['es', 'en']:
                    cache.set(cache_key, lang, CACHE_TIMEOUT)
                    return lang
        lang = detect(text[:1000])
        if lang in ['es', 'en']:
            cache.set(cache_key, lang, CACHE_TIMEOUT)
            return lang
        return default
    except (LangDetectException, LDE, Exception) as e:
        logger.warning(f"Language detection error: {e}")
        return default

@lru_cache(maxsize=4096)
def normalize_text(text: str, lang: str = None) -> str:
    """Normalize text by removing accents and converting to lowercase."""
    if not text:
        return ""
    cache_key = f'norm_text:{hash(text[:100])}:{lang}'
    cached = cache.get(cache_key)
    if cached:
        return cached
    try:
        normalized = ''.join(
            char for char in unicodedata.normalize('NFD', str(text))
            if unicodedata.category(char) != 'Mn'
        ).lower().strip()
        if lang == 'es':
            normalized = normalized.replace('ñ', 'n')
        cache.set(cache_key, normalized, CACHE_TIMEOUT)
        return normalized
    except Exception as e:
        logger.error(f"Error normalizing text: {e}")
        return str(text).lower().strip()

def load_division_skills() -> Dict[str, List[str]]:
    """Load division skills once and store in global cache."""
    global DIVISION_SKILLS_CACHE
    if DIVISION_SKILLS_CACHE is None:
        try:
            divisions = Division.objects.all()
            DIVISION_SKILLS_CACHE = {
                division.name: list(Skill.objects.filter(division=division).values_list('name', flat=True))
                for division in divisions
            }
        except Exception as e:
            logger.error(f"Error loading division skills: {e}")
            DIVISION_SKILLS_CACHE = {"finance": ["budgeting"], "tech": ["programming"]}
    return DIVISION_SKILLS_CACHE

class IMAPCVProcessor:
    def __init__(self, business_unit: BusinessUnit, batch_size: int = 10, sleep_time: float = 2.0):
        self.business_unit = business_unit
        self.config = self._load_config(business_unit)
        self.parser = CVParser(business_unit)
        self.stats = {"processed": 0, "created": 0, "updated": 0, "errors": 0}
        self.FOLDER_CONFIG = FOLDER_CONFIG
        self.batch_size = batch_size
        self.sleep_time = sleep_time

    def _load_config(self, business_unit: BusinessUnit) -> Dict:
        try:
            config = ConfiguracionBU.objects.get(business_unit=business_unit)
            return {
                'server': 'mail.huntred.com',
                'port': 993,
                'username': config.smtp_username,
                'password': config.smtp_password,
            }
        except ConfiguracionBU.DoesNotExist:
            raise ValueError(f"No configuration for business unit: {business_unit.name}")

    async def _connect_imap(self, config: Dict) -> Optional[aioimaplib.IMAP4_SSL]:
        try:
            client = aioimaplib.IMAP4_SSL(config['server'], config['port'])
            await client.wait_hello_from_server()
            await client.login(config['username'], config['password'])
            if not await self._verify_folders(client):
                raise ValueError("IMAP folders not configured correctly")
            logger.info(f"Connected to IMAP server: {config['server']}")
            return client
        except Exception as e:
            logger.error(f"Error connecting to IMAP server: {e}")
            return None

    async def _verify_folders(self, mail) -> bool:
        for folder_key, folder_name in self.FOLDER_CONFIG.items():
            try:
                resp, folder_list = await mail.list(pattern=folder_name)
                if not folder_list:
                    logger.error(f"Folder not found: {folder_name}")
                    return False
            except Exception as e:
                logger.error(f"Error verifying folder {folder_name}: {e}")
                return False
        return True

    async def _move_email(self, mail, msg_id: str, dest_folder: str):
        try:
            await mail.copy(msg_id, dest_folder)
            await mail.store(msg_id, '+FLAGS', '\\Deleted')
            logger.info(f"Email {msg_id} moved to {dest_folder}")
        except Exception as e:
            logger.error(f"Error moving email {msg_id} to {dest_folder}: {e}")

    def _update_stats(self, result: Dict):
        if result.get("status") == "created":
            self.stats["created"] += 1
        elif result.get("status") == "updated":
            self.stats["updated"] += 1
        self.stats["processed"] += 1

    async def _process_single_email(self, mail, email_id: str):
        try:
            resp, data = await mail.fetch(email_id, "(RFC822)")
            if resp != "OK":
                raise ValueError(f"Error fetching email {email_id}")
            message = email.message_from_bytes(data[0][1])
            attachments = self.parser.extract_attachments(message)
            if not attachments:
                logger.warning(f"Email {email_id} has no valid attachments")
                await self._move_email(mail, email_id, self.FOLDER_CONFIG['error_folder'])
                self.stats["errors"] += 1
                return
            with ExitStack() as stack:
                for attachment in attachments:
                    temp_file = stack.enter_context(NamedTemporaryFile(delete=False))
                    temp_path = Path(temp_file.name)
                    temp_path.write_bytes(attachment['content'])
                    text = self.parser.extract_text_from_file(temp_path)
                    if text:
                        parsed_data = await self.parser.parse(text)
                        if parsed_data:
                            email_addr = parsed_data.get("email", "")
                            phone = parsed_data.get("phone", "")
                            candidate = await sync_to_async(Person.objects.filter)(email=email_addr).first() or \
                                        await sync_to_async(Person.objects.filter)(phone=phone).first()
                            if candidate:
                                await self.parser._update_candidate(candidate, parsed_data, temp_path)
                                self._update_stats({"status": "updated"})
                            else:
                                await self.parser._create_new_candidate(parsed_data, temp_path)
                                self._update_stats({"status": "created"})
            await self._move_email(mail, email_id, self.FOLDER_CONFIG['parsed_folder'])
        except Exception as e:
            logger.error(f"Error processing email {email_id}: {e}")
            await self._move_email(mail, email_id, self.FOLDER_CONFIG['error_folder'])
            self.stats["errors"] += 1
            if self.business_unit.admin_email:
                email_service = EmailService(self.business_unit)
                await email_service.send_email(
                    subject=f"Error en CV Parser: {email_id}",
                    to_email=self.business_unit.admin_email,
                    body=f"Error processing CV in email {email_id}: {str(e)}",
                    from_email="noreply@huntred.com"
                )

    async def process_emails(self):
        mail = await self._connect_imap(self.config)
        if not mail:
            return
        try:
            resp, _ = await mail.select(self.FOLDER_CONFIG['cv_folder'])
            if resp != "OK":
                raise ValueError(f"Error selecting {self.FOLDER_CONFIG['cv_folder']}")
            resp, messages = await mail.search("ALL")
            if resp != "OK":
                raise ValueError("Error searching messages")
            email_ids = messages[0].split()
            logger.info(f"Total emails to process: {len(email_ids)}")
            for i in range(0, len(email_ids), self.batch_size):
                batch = email_ids[i:i + self.batch_size]
                logger.info(f"Processing batch of {len(batch)} emails")
                await asyncio.gather(*[self._process_single_email(mail, email_id) for email_id in batch])
                await asyncio.sleep(self.sleep_time)
            await mail.expunge()
            logger.info("Expunged deleted emails")
        finally:
            await mail.logout()
            logger.info("Disconnected from IMAP server")
            await self._generate_summary_and_send_report(**self.stats)

    async def _generate_summary_and_send_report(self, processed: int, created: int, updated: int, errors: int):
        admin_email = self.business_unit.admin_email
        if not admin_email:
            logger.warning("Admin email not configured")
            return
        summary = f"""
        <h2>CV Processing Summary for {self.business_unit.name}:</h2>
        <ul>
            <li>Total emails processed: {processed}</li>
            <li>New candidates created: {created}</li>
            <li>Candidates updated: {updated}</li>
            <li>Errors encountered: {errors}</li>
        </ul>
        """
        logger.info(f"Summary generated:\n{summary}")
        email_service = EmailService(self.business_unit)
        await email_service.send_email(
            subject=f"CV Processing Summary - {self.business_unit.name}",
            to_email=admin_email,
            body=summary,
            from_email="noreply@huntred.com"
        )

class CVParser:
    """Parser especializado para CVs."""
    
    def __init__(self, business_unit: BusinessUnit, max_workers: int = None):
        self.business_unit = business_unit
        self.max_workers = max_workers or MAX_WORKERS
        self.executor = ThreadPoolExecutor(max_workers=self.max_workers)
        self.nlp = get_nlp()
        self.matcher = Matcher(self.nlp.vocab) if self.nlp else None
        self._init_skill_patterns()
        self._init_legacy()
        self._init_business_unit_skills()
        self.text_processor = TextProcessor()
        self.validator = DocumentValidator()
        self.scraping_utils = PlaywrightAntiDeteccion()

    def _init_skill_patterns(self):
        if not self.matcher:
            return
        patterns = [
            [{"LOWER": "python"}],
            [{"LOWER": "java"}],
            [{"LOWER": "javascript"}],
            [{"LOWER": "machine"}, {"LOWER": "learning"}],
            [{"LOWER": "data"}, {"LOWER": "science"}],
            [{"LOWER": "artificial"}, {"LOWER": "intelligence"}],
            [{"LOWER": "cloud"}, {"LOWER": "computing"}],
            [{"LOWER": "devops"}],
            [{"LOWER": "agile"}],
            [{"LOWER": "scrum"}],
        ]
        self.matcher.add("SKILLS", patterns)

    def _init_legacy(self):
        self.DIVISION_SKILLS = load_division_skills()
        self.email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
        self.phone_pattern = r'\b\d{10}\b'

    def _init_business_unit_skills(self):
        self.HUNTRED_TECH_SKILLS = {
            'python', 'django', 'fastapi', 'sql', 'nosql', 'aws', 'gcp', 'azure',
            'devops', 'ci/cd', 'kubernetes', 'docker', 'terraform', 'ansible',
            'machine_learning', 'deep_learning', 'nlp', 'computer_vision'
        }
        self.HUNTRED_SOFT_SKILLS = {'leadership', 'teamwork', 'communication'}
        self.HUNTRED_TOOLS = {'git', 'jira', 'confluence'}
        self.HUNTRED_CERTS = {'aws certified', 'pmp', 'scrum master'}
        self.HUNTU_TECH_SKILLS = {'javascript', 'react', 'node.js'}
        self.HUNTU_SOFT_SKILLS = {'adaptability', 'problem-solving'}
        self.HUNTU_TOOLS = {'vscode', 'slack'}
        self.HUNTU_CERTS = {'react certified', 'node.js certified'}
        self.SEXSI_TECH_SKILLS = {'communication', 'psychology'}
        self.SEXSI_SOFT_SKILLS = {'empathy', 'active listening'}
        self.SEXSI_TOOLS = {'zoom', 'teams'}
        self.SEXSI_CERTS = {'coaching certified', 'hr certified'}

    def __del__(self):
        self.executor.shutdown(wait=True)

    async def parse_resume(self, file_content: bytes, file_extension: str = None) -> Dict:
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            self.executor,
            self._parse_resume_sync,
            file_content,
            file_extension
        )

    def _parse_resume_sync(self, file_content: bytes, file_extension: str = None) -> Dict:
        cpu_before = psutil.cpu_percent()
        try:
            if HAS_RESUME_PARSER and file_extension in ['pdf', 'docx', 'doc']:
                with NamedTemporaryFile(suffix=f'.{file_extension}', delete=False) as temp_file:
                    temp_path = temp_file.name
                    temp_file.write(file_content)
                try:
                    data = resumeparse.read_file(temp_path)
                    data.pop("degree", None)  # Skip degree parsing
                    result = self._process_parsed_resume(data)
                    logger.info("Parsed resume with %d skills using resume_parser", len(result.get("skills", [])))
                    return result
                except Exception as e:
                    logger.warning(f"resume_parser failed: {e}")
                finally:
                    try:
                        Path(temp_path).unlink()
                    except Exception:
                        pass
            text = self._extract_text(file_content, file_extension)
            if not text:
                return {"error": "No text extracted", "status": "error"}
            lang = detect_language(text)
            normalized_text = normalize_text(text, lang)
            email = self._extract_email(normalized_text)
            phone = self._extract_phone(normalized_text)
            result = {
                "email": email,
                "phone": phone,
                "language": lang,
                "text_length": len(normalized_text),
                "status": "success"
            }
            logger.info("Parsed resume with fallback method, text length: %d", len(normalized_text))
            return result
        except Exception as e:
            logger.error(f"Error parsing resume: {e}", exc_info=True)
            return {"error": str(e), "status": "error"}
        finally:
            cpu_after = psutil.cpu_percent()
            logger.info("CPU usage during resume parsing: %s%%", cpu_after - cpu_before)
            gc.collect()

    async def parse_resume_batch(self, files: List[Tuple[bytes, str]]) -> List[Dict]:
        results = []
        chunk_size = self._calculate_optimal_chunk_size(files)
        for i in range(0, len(files), chunk_size):
            chunk = files[i:i + chunk_size]
            try:
                chunk_results = await self._process_chunk(chunk)
                results.extend(chunk_results)
                if await self._check_system_health():
                    await asyncio.sleep(2.0)
            except Exception as e:
                logger.error(f"Error processing chunk: {e}")
        return results

    def _calculate_optimal_chunk_size(self, files: List[Tuple[bytes, str]]) -> int:
        if not files:
            return 1
        total_size = sum(len(content) for content, _ in files)
        avg_size = total_size / len(files)
        if avg_size > 1000000:
            return 1
        elif avg_size > 500000:
            return 2
        elif avg_size > 100000:
            return 5
        else:
            return 10

    async def _process_chunk(self, chunk: List[Tuple[bytes, str]]) -> List[Dict]:
        tasks = [asyncio.create_task(self._process_single_file(content, extension)) for content, extension in chunk]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        valid_results = []
        for result in results:
            if isinstance(result, Exception):
                logger.error(f"Error processing file: {result}")
                continue
            if result:
                valid_results.append(result)
        return valid_results

    async def _process_single_file(self, content: bytes, extension: str) -> Optional[Dict]:
        try:
            with NamedTemporaryFile(suffix=extension, delete=False) as temp_file:
                temp_path = Path(temp_file.name)
                temp_path.write_bytes(content)
                text = self.extract_text_from_file(temp_path)
                if not text:
                    raise ValueError("No text extracted")
                doc = self.nlp(text)
                entities = self._extract_entities(text)
                validated_entities = self._validate_entities(entities)
                experience = self._process_experience(doc, self.business_unit)
                education = self._process_education(doc, self.business_unit)
                experience_level = self._analyze_experience_level(doc)
                sentiment_analysis = self._analyze_sentiment_and_tone(text)
                welcome_message = self._generate_welcome_message(
                    self.business_unit, experience_level, sentiment_analysis
                )
                return {
                    "text": text,
                    "entities": validated_entities,
                    "experience": experience,
                    "education": education,
                    "experience_level": experience_level,
                    "sentiment_analysis": sentiment_analysis,
                    "welcome_message": welcome_message,
                    "metadata": {
                        "processed_at": datetime.now().isoformat(),
                        "file_type": extension,
                        "file_size": len(content)
                    }
                }
        except Exception as e:
            logger.error(f"Error processing file: {e}")
            return None
        finally:
            try:
                if temp_path.exists():
                    temp_path.unlink()
            except Exception:
                logger.warning(f"Error deleting temp file: {e}")

    async def _check_system_health(self) -> bool:
        try:
            process = psutil.Process()
            memory_info = process.memory_info()
            if memory_info.rss > 500 * 1024 * 1024:  # > 500MB
                logger.warning("High memory usage detected")
                gc.collect()
                return True
            cpu_percent = process.cpu_percent(interval=0.1)
            if cpu_percent > 80:
                logger.warning("High CPU usage detected")
                return True
            return False
        except Exception as e:
            logger.error(f"Error checking system health: {e}")
            return False

    async def _handle_error(self, error: Exception, context: str) -> None:
        error_type = type(error).__name__
        error_msg = str(error)
        logger.error(f"Error in {context}: {error_type} - {error_msg}")
        if isinstance(error, (MemoryError, OSError)):
            await self._send_error_notification(error_type, error_msg, context)
        if isinstance(error, MemoryError):
            self.max_workers = max(1, self.max_workers - 1)
        elif isinstance(error, TimeoutError):
            self.sleep_time *= 1.5

    async def _send_error_notification(self, error_type: str, error_msg: str, context: str) -> None:
        try:
            subject = f"Critical Error in CV Parser: {error_type}"
            body = f"""
            Context: {context}
            Error: {error_msg}
            Timestamp: {datetime.now().isoformat()}
            """
            await sync_to_async(send_mail)(
                subject=subject,
                message=body,
                from_email="noreply@huntred.com",
                recipient_list=["admin@huntred.com"],
                fail_silently=True
            )
        except Exception as e:
            logger.error(f"Error sending notification: {e}")

    def _extract_text(self, file_content: bytes, file_extension: str) -> str:
        try:
            if file_extension == 'pdf':
                return self._extract_text_from_pdf(file_content)
            elif file_extension in ['docx', 'doc']:
                return self._extract_text_from_docx(file_content)
            elif file_extension in ['txt', 'md']:
                return file_content.decode('utf-8', errors='ignore')
            else:
                logger.warning(f"Unsupported file extension: {file_extension}")
                return ""
        except Exception as e:
            logger.error(f"Error extracting text: {e}")
            return ""

    def _extract_email(self, text: str) -> str:
        email_match = re.search(r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+', text)
        return email_match.group(0) if email_match else ""

    def _extract_phone(self, text: str) -> str:
        phone_match = re.search(
            r'(\+\d{1,3}[-.\s]?)?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,9}',
            text
        )
        return phone_match.group(0) if phone_match else ""

    def _process_parsed_resume(self, parsed: Dict) -> Dict:
        try:
            result = {
                'name': parsed.get('name', '').strip(),
                'email': parsed.get('email', '').lower().strip(),
                'phone': parsed.get('phone', '').strip(),
                'skills': [s.strip() for s in parsed.get('skills', []) if s.strip()],
                'education': [e.strip() for e in parsed.get('education', []) if e.strip()],
                'experience': parsed.get('experience', []),
                'languages': [l.strip() for l in parsed.get('languages', []) if l.strip()],
                'raw': parsed
            }
            text = ' '.join([
                result['name'],
                ' '.join(result['skills']),
                ' '.join(result['education'])
            ])
            result['language'] = detect_language(text)
            return result
        except Exception as e:
            logger.error(f"Error processing parsed resume: {e}", exc_info=True)
            return {"error": str(e), "status": "error", "raw": parsed}

    async def parse(self, text: str) -> Dict:
        try:
            detected_lang = detect_language(text)
            logger.info(f"Detected language: {detected_lang}")
            doc = self.nlp(text)
            skills = self._extract_skills(doc, self.business_unit)
            entities = self._extract_entities(text)
            experience = self._process_experience(doc, self.business_unit)
            education = self._process_education(doc, self.business_unit)
            experience_level = self._analyze_experience_level(doc)
            sentiment_analysis = self._analyze_sentiment_and_tone(text)
            welcome_message = self._generate_welcome_message(self.business_unit, experience_level, sentiment_analysis)
            return {
                "email": entities["email"],
                "phone": entities["phone"],
                "skills": skills,
                "experience": experience,
                "education": education,
                "sentiment_analysis": sentiment_analysis,
                "experience_level": experience_level,
                "welcome_message": welcome_message
            }
        except Exception as e:
            logger.error(f"Error parsing CV text: {e}")
            return {}

    def extract_attachments(self, message) -> List[Dict]:
        attachments = []
        for part in message.walk():
            if part.get_content_maintype() == 'multipart':
                continue
            if part.get('Content-Disposition') is None:
                continue
            filename = part.get_filename()
            if filename:
                content = part.get_payload(decode=True)
                attachments.append({'filename': filename, 'content': content})
        logger.info(f"Found {len(attachments)} attachments")
        return attachments

    def _extract_skills(self, analysis: Doc, business_unit: BusinessUnit) -> Dict:
        skills = {"technical": [], "soft": [], "tools": [], "certifications": []}
        if not self.matcher:
            return skills
        matches = self.matcher(analysis)
        skills["technical"] = [analysis[start:end].text for _, start, end in matches]
        if business_unit.name.lower() == 'huntred':
            return self._filter_huntred_skills(skills)
        elif business_unit.name.lower() == 'huntu':
            return self._filter_huntu_skills(skills)
        elif business_unit.name.lower() == 'sexsi':
            return self._filter_sexsi_skills(skills)
        return skills

    def _filter_huntred_skills(self, skills: Dict) -> Dict:
        return {
            "technical": [s for s in skills["technical"] if s.lower() in self.HUNTRED_TECH_SKILLS],
            "soft": [s for s in skills["soft"] if s.lower() in self.HUNTRED_SOFT_SKILLS],
            "tools": [s for s in skills["tools"] if s.lower() in self.HUNTRED_TOOLS],
            "certifications": [s for s in skills["certifications"] if s.lower() in self.HUNTRED_CERTS]
        }

    def _filter_huntu_skills(self, skills: Dict) -> Dict:
        return {
            "technical": [s for s in skills["technical"] if s.lower() in self.HUNTU_TECH_SKILLS],
            "soft": [s for s in skills["soft"] if s.lower() in self.HUNTU_SOFT_SKILLS],
            "tools": [s for s in skills["tools"] if s.lower() in self.HUNTU_TOOLS],
            "certifications": [s for s in skills["certifications"] if s.lower() in self.HUNTU_CERTS]
        }

    def _filter_sexsi_skills(self, skills: Dict) -> Dict:
        return {
            "technical": [s for s in skills["technical"] if s.lower() in self.SEXSI_TECH_SKILLS],
            "soft": [s for s in skills["soft"] if s.lower() in self.SEXSI_SOFT_SKILLS],
            "tools": [s for s in skills["tools"] if s.lower() in self.SEXSI_TOOLS],
            "certifications": [s for s in skills["certifications"] if s.lower() in self.SEXSI_CERTS]
        }

    def _extract_entities(self, text: str) -> Dict:
        entities = {
            "name": [], "email": [], "phone": [], "organization": [], "location": [],
            "skills": [], "education": [], "experience": [], "certifications": [], "languages": []
        }
        doc = self.nlp(text)
        for ent in doc.ents:
            if ent.label_ == "PER":
                entities["name"].append(ent.text)
            elif ent.label_ == "ORG":
                entities["organization"].append(ent.text)
            elif ent.label_ == "LOC":
                entities["location"].append(ent.text)
            elif ent.label_ == "DATE":
                if "año" in ent.text.lower() or "year" in ent.text.lower():
                    entities["experience"].append(ent.text)
        entities["email"] = re.findall(self.email_pattern, text)
        entities["phone"] = re.findall(self.phone_pattern, text)
        if self.matcher:
            matches = self.matcher(doc)
            entities["skills"] = [doc[start:end].text for _, start, end in matches]
        cert_pattern = r'(?:AWS|Azure|GCP|Oracle|Microsoft|Cisco|CompTIA|PMP|ITIL|Scrum|Agile)[^.,\n]*'
        entities["certifications"] = re.findall(cert_pattern, text, re.IGNORECASE)
        lang_pattern = r'(?:Inglés|Español|Francés|Alemán|Italiano|Portugués|Chino|Japonés|Ruso)[^.,\n]*'
        entities["languages"] = re.findall(lang_pattern, text, re.IGNORECASE)
        education_pattern = r'(?:Bachelor|Master|PhD|B\.S\.|M\.S\.|B\.A\.|M\.A\.|Licenciatura|Maestría|Doctorado|Ingeniería|Grado|Diplomatura)[^.,\n]*'
        entities["education"] = re.findall(education_pattern, text, re.IGNORECASE)
        experience_pattern = r'(?:\d+\s*(?:años|years|yr|yrs)\s+de\s+experiencia|experience|experiencia)[^.,\n]*'
        entities["experience"] = re.findall(experience_pattern, text, re.IGNORECASE)
        return entities

    def _validate_entities(self, entities: Dict) -> Dict:
        validated = {}
        validated["email"] = next(
            (email for email in entities["email"] if re.match(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$', email)),
            ""
        )
        validated["phone"] = next(
            (phone for phone in entities["phone"] if re.match(r'^\+?\d{8,15}$', phone.replace(' ', '').replace('-', ''))),
            ""
        )
        names = [name for name in entities["name"] if len(name.split()) >= 2]
        validated["name"] = names[0] if names else ""
        validated["organizations"] = list(set(org.strip() for org in entities["organization"] if org.strip() and len(org) > 2))
        validated["locations"] = list(set(loc.strip() for loc in entities["location"] if loc.strip() and len(loc) > 2))
        validated["skills"] = list(set(skill.strip().lower() for skill in entities["skills"] if skill.strip() and len(skill) > 2))
        validated["certifications"] = list(set(cert.strip() for cert in entities["certifications"] if cert.strip() and len(cert) > 2))
        validated["languages"] = list(set(lang.strip() for lang in entities["languages"] if lang.strip() and len(lang) > 2))
        validated["education"] = list(set(edu.strip() for edu in entities["education"] if edu.strip() and len(edu) > 2))
        validated["experience"] = list(set(exp.strip() for exp in entities["experience"] if exp.strip() and len(exp) > 2))
        return validated

    def _process_experience(self, analysis: Doc, business_unit: BusinessUnit) -> List:
        experience = []
        for sent in analysis.sents:
            if any(keyword in sent.text.lower() for keyword in ["experience", "experiencia", "worked", "trabajó"]):
                experience.append({"text": sent.text.strip()})
        if business_unit.name.lower() == 'huntred':
            return self._process_huntred_experience(experience)
        elif business_unit.name.lower() == 'huntu':
            return self._process_huntu_experience(experience)
        elif business_unit.name.lower() == 'sexsi':
            return self._process_sexsi_experience(experience)
        return experience

    def _process_huntred_experience(self, experience: List) -> List:
        processed = []
        for exp in experience:
            text = exp["text"].lower()
            if any(skill in text for skill in self.HUNTRED_TECH_SKILLS):
                processed.append(exp)
        return processed

    def _process_huntu_experience(self, experience: List) -> List:
        processed = []
        for exp in experience:
            text = exp["text"].lower()
            if any(skill in text for skill in self.HUNTU_TECH_SKILLS):
                processed.append(exp)
        return processed

    def _process_sexsi_experience(self, experience: List) -> List:
        processed = []
        for exp in experience:
            text = exp["text"].lower()
            if any(skill in text for skill in self.SEXSI_TECH_SKILLS):
                processed.append(exp)
        return processed

    def _process_education(self, analysis: Doc, business_unit: BusinessUnit) -> List:
        education = []
        for sent in analysis.sents:
            if any(keyword in sent.text.lower() for keyword in ["bachelor", "master", "phd", "licenciatura", "maestría"]):
                education.append({"text": sent.text.strip()})
        if business_unit.name.lower() == 'huntred':
            return self._process_huntred_education(education)
        elif business_unit.name.lower() == 'huntu':
            return self._process_huntu_education(education)
        elif business_unit.name.lower() == 'sexsi':
            return self._process_sexsi_education(education)
        return education

    def _process_huntred_education(self, education: List) -> List:
        processed = []
        for edu in education:
            text = edu["text"].lower()
            if any(word in text for word in ["computer", "software", "engineering", "technology", "data", "analytics"]):
                processed.append(edu)
        return processed

    def _process_huntu_education(self, education: List) -> List:
        processed = []
        for edu in education:
            text = edu["text"].lower()
            if any(word in text for word in ["computer", "software", "engineering", "technology", "data", "analytics"]):
                processed.append(edu)
        return processed

    def _process_sexsi_education(self, education: List) -> List:
        processed = []
        for edu in education:
            text = edu["text"].lower()
            if any(word in text for word in ["communication", "psychology", "sociology", "humanities"]):
                processed.append(edu)
        return processed

    def _analyze_experience_level(self, analysis: Doc) -> Dict:
        years = 0
        for sent in analysis.sents:
            if any(keyword in sent.text.lower() for keyword in ["year", "año", "years", "años"]):
                match = re.search(r'(\d+)\s*(?:year|years|año|años)', sent.text, re.IGNORECASE)
                if match:
                    years += int(match.group(1))
        if years >= 10:
            return {"level": "expert", "years": years}
        elif years >= 5:
            return {"level": "senior", "years": years}
        elif years >= 2:
            return {"level": "mid", "years": years}
        else:
            return {"level": "junior", "years": years}

    def _analyze_sentiment_and_tone(self, text: str) -> Dict:
        try:
            doc = self.nlp(text)
            sentiment_score = 0
            positive_words = {'excelente', 'experto', 'experiencia', 'habilidad', 'logro', 'éxito', 'innovador'}
            negative_words = {'básico', 'limitado', 'poco', 'mínimo', 'inicial'}
            for token in doc:
                if token.text.lower() in positive_words:
                    sentiment_score += 1
                elif token.text.lower() in negative_words:
                    sentiment_score -= 1
            sentiment = (
                "muy positivo" if sentiment_score > 2 else
                "positivo" if sentiment_score > 0 else
                "muy negativo" if sentiment_score < -2 else
                "negativo" if sentiment_score < 0 else
                "neutral"
            )
            tone = {"profesional": 0, "técnico": 0, "creativo": 0, "formal": 0}
            tone_keywords = {
                "profesional": ['experiencia', 'responsabilidad', 'liderazgo', 'gestión', 'equipo'],
                "técnico": ['tecnología', 'desarrollo', 'implementación', 'optimización', 'arquitectura'],
                "creativo": ['innovación', 'diseño', 'creatividad', 'solución', 'idea'],
                "formal": ['objetivo', 'meta', 'resultado', 'proceso', 'sistema']
            }
            for tone_type, keywords in tone_keywords.items():
                for keyword in keywords:
                    tone[tone_type] += len(re.findall(r'\b' + keyword + r'\b', text.lower()))
            dominant_tone = max(tone.items(), key=lambda x: x[1])[0]
            confidence_indicators = [
                'experto', 'experiencia', 'dominio', 'conocimiento', 'habilidad',
                'capacidad', 'aptitud', 'competencia', 'maestría', 'pericia'
            ]
            confidence_score = sum(1 for word in confidence_indicators if word in text.lower())
            confidence_level = "alto" if confidence_score > 5 else "medio" if confidence_score > 2 else "bajo"
            return {
                "sentiment": {"score": sentiment_score, "label": sentiment},
                "tone": {"analysis": tone, "dominant": dominant_tone},
                "confidence": {"score": confidence_score, "level": confidence_level}
            }
        except Exception as e:
            logger.error(f"Error in sentiment analysis: {e}")
            return {
                "sentiment": {"score": 0, "label": "neutral"},
                "tone": {"analysis": {}, "dominant": "neutral"},
                "confidence": {"score": 0, "level": "bajo"}
            }

    def _generate_welcome_message(self, business_unit: BusinessUnit, experience_level: Dict, sentiment_analysis: Dict = None) -> str:
        bu_name = business_unit.name.lower()
        experience_years = experience_level.get('years', 0)
        experience_label = experience_level.get('level', 'junior')
        base_messages = {
            'huntred': f"¡Bienvenido/a a HuntRED! 💼 Con {experience_years} años de experiencia, ",
            'huntu': f"¡Bienvenido/a a Huntu! 🏆 Con {experience_years} años de experiencia, ",
            'sexsi': f"¡Bienvenido/a a SEXSI! 📜 Con {experience_years} años de experiencia, "
        }
        experience_messages = {
            'expert': "tu perfil destaca por su amplia experiencia y conocimientos avanzados. ",
            'senior': "tienes un perfil sólido y bien establecido. ",
            'mid': "tienes un buen nivel de experiencia. ",
            'junior': "estás comenzando tu carrera profesional. "
        }
        tone_messages = {
            'profesional': "Tu enfoque profesional es muy valorado. ",
            'técnico': "Tu perfil técnico es muy interesante. ",
            'creativo': "Tu enfoque creativo es muy apreciado. ",
            'formal': "Tu perfil formal es muy adecuado. "
        }
        confidence_messages = {
            'alto': "Tu perfil muestra gran seguridad y dominio. ",
            'medio': "Tu perfil muestra buena confianza. ",
            'bajo': "Tu perfil muestra potencial de crecimiento. "
        }
        message = base_messages.get(bu_name, "¡Bienvenido/a a Grupo huntRED! ")
        message += experience_messages.get(experience_label, "")
        if sentiment_analysis:
            message += tone_messages.get(sentiment_analysis['tone']['dominant'], "")
            message += confidence_messages.get(sentiment_analysis['confidence']['level'], "")
        message += "Por favor, responde con el código de verificación que recibirá en su email."
        return message

    def extract_text_from_file(self, file_path: Path) -> Optional[str]:
        try:
            if file_path.suffix.lower() == '.pdf':
                return self._extract_text_from_pdf(file_path)
            elif file_path.suffix.lower() in ['.doc', '.docx']:
                return self._extract_text_from_docx(file_path)
            elif file_path.suffix.lower() == '.html':
                return self._extract_text_from_html(file_path)
            elif file_path.suffix.lower() in ['.jpg', '.jpeg', '.png']:
                return self._extract_text_from_image(file_path)
            else:
                logger.warning(f"Unsupported format: {file_path}")
                return None
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return None

    def _extract_text_from_pdf(self, file_path: Union[Path, bytes]) -> str:
        text = ""
        try:
            if isinstance(file_path, bytes):
                doc = fitz.open(stream=file_path, filetype="pdf")
            else:
                doc = fitz.open(str(file_path))
            for page in doc:
                text += page.get_text("text", sort=True)
                tables = page.find_tables()
                if tables:
                    for table in tables:
                        text += "\n" + table.to_text()
            doc.close()
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            try:
                if isinstance(file_path, bytes):
                    with io.BytesIO(file_path) as f:
                        with pdfplumber.open(f) as pdf:
                            text = ""
                            for page in pdf.pages:
                                text += page.extract_text() or ""
                                tables = page.extract_tables()
                                if tables:
                                    for table in tables:
                                        text += "\n" + "\n".join([" | ".join(row) for row in table])
                            return text.strip()
                else:
                    with pdfplumber.open(str(file_path)) as pdf:
                        text = ""
                        for page in pdf.pages:
                            text += page.extract_text() or ""
                            tables = page.extract_tables()
                            if tables:
                                for table in tables:
                                    text += "\n" + "\n".join([" | ".join(row) for row in table])
                        return text.strip()
            except Exception as e2:
                logger.error(f"PDF fallback error: {e2}")
                return ""

    def _extract_text_from_docx(self, file_path: Union[Path, bytes]) -> str:
        try:
            if isinstance(file_path, bytes):
                with io.BytesIO(file_path) as f:
                    doc = Document(f)
            else:
                doc = Document(str(file_path))
            return "\n".join(para.text for para in doc.paragraphs)
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            return ""

    def _extract_text_from_html(self, file_path: Path) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        return trafilatura.extract(html_content) or ""

    def _extract_text_from_image(self, file_path: Path) -> str:
        try:
            image = Image.open(str(file_path)).convert("L")
            text = pytesseract.image_to_string(image)
            return text
        except Exception as e:
            logger.error(f"Error in OCR: {e}")
            return ""

    async def _update_candidate(self, candidate: Person, parsed_data: Dict, file_path: Path):
        candidate.cv_file = str(file_path)
        candidate.cv_analysis = parsed_data
        candidate.cv_parsed = True
        candidate.metadata.update({
            "last_cv_update": now().isoformat(),
            "skills": parsed_data["skills"],
            "experience": parsed_data["experience"],
            "education": parsed_data["education"],
            "sentiment": parsed_data["sentiment"]
        })
        await sync_to_async(candidate.save)()
        logger.info(f"Updated profile: {candidate.nombre} {candidate.apellido_paterno}")

    async def _create_new_candidate(self, parsed_data: Dict, file_path: Path):
        """Crea un nuevo candidato de forma asíncrona."""
        candidate = await sync_to_async(Person.objects.create)(
            nombre=parsed_data.get("nombre", "Unknown"),
            apellido_paterno=parsed_data.get("apellido_paterno", ""),
            email=parsed_data.get("email", ""),
            phone=parsed_data.get("phone", ""),
            cv_file=str(file_path),
            cv_parsed=True,
            cv_analysis=parsed_data,
            metadata={
                "last_cv_update": now().isoformat(),
                "skills": parsed_data["skills"],
                "experience": parsed_data["experience"],
                "education": parsed_data["education"],
                "sentiment": parsed_data["sentiment"]
            }
        )
        logger.info(f"✅ Nuevo perfil creado: {candidate.nombre} {candidate.apellido_paterno}")

def parse_document(file_path: str, business_unit_name: str) -> Dict:
    """
    Parse a CV document from a file path.
    
    Example usage:
    ```python
    result = parse_document("path/to/resume.pdf", "huntRED")
    print(f"Extracted email: {result.get('email')}")
    print(f"Skills: {', '.join(result.get('skills', []))}")
    ```
    
    Args:
        file_path: Path to the CV file
        business_unit_name: Name of the business unit
        
    Returns:
        Dict: Parsed CV data or error information
    """
    try:
        business_unit = BusinessUnit.objects.get(name=business_unit_name)
        parser = CVParser(business_unit)
        
        with open(file_path, 'rb') as f:
            file_content = f.read()
        
        file_extension = Path(file_path).suffix[1:].lower()
        return parser._parse_resume_sync(file_content, file_extension)
        
    except Exception as e:
        logger.error(f"Error parsing document {file_path}: {e}", exc_info=True)
        return {"error": str(e), "status": "error"}

async def parse_document_async(file_path: str, business_unit_name: str) -> Dict:
    """
    Asynchronously parse a CV document from a file path.
    
    Example usage:
    ```python
    result = await parse_document_async("path/to/resume.pdf", "huntRED")
    print(f"Extracted email: {result.get('email')}")
    print(f"Skills: {', '.join(result.get('skills', []))}")
    ```
    
    Args:
        file_path: Path to the CV file
        business_unit_name: Name of the business unit
        
    Returns:
        Dict: Parsed CV data or error information
    """
    try:
        business_unit = await sync_to_async(BusinessUnit.objects.get)(name=business_unit_name)
        parser = CVParser(business_unit)
        
        with open(file_path, 'rb') as f:
            file_content = f.read()
        
        file_extension = Path(file_path).suffix[1:].lower()
        return await parser.parse_resume(file_content, file_extension)
        
    except Exception as e:
        logger.error(f"Error parsing document {file_path}: {e}", exc_info=True)
        return {"error": str(e), "status": "error"}

def parse_job_listing(content, source_url, source_type="web"):
    """Parse job listing content from web or email sources to extract structured information."""
    try:
        soup = BeautifulSoup(content, "html.parser")
        job_info = {}
        
        if source_type == "web":
            title_elem = soup.find(["h1", "h2", "h3"], class_=["title", "job-title", "header"]) or soup.find(["h1", "h2", "h3"])
            company_elem = soup.find(["span", "div"], class_=["company", "employer", "organization"])
            desc_elem = soup.find(["div", "section"], class_=["description", "summary", "details", "content"])
            url_elem = soup.find("a", href=True, text=re.compile(r"apply|postular|solicitar", re.I))
            
            job_info["title"] = title_elem.get_text(strip=True) if title_elem else "Untitled Job"
            job_info["company"] = company_elem.get_text(strip=True) if company_elem else "Unknown Company"
            job_info["description"] = desc_elem.get_text(strip=True)[:1000] if desc_elem else "No description available"
            job_info["url"] = urljoin(source_url, url_elem["href"]) if url_elem else source_url
            job_info["source"] = urlparse(source_url).netloc
        
        elif source_type == "email":
            job_info["title"] = soup.title.get_text(strip=True) if soup.title else "Untitled Email Job"
            company_elem = soup.find(text=re.compile(r"from|de|empresa|company", re.I))
            job_info["company"] = company_elem.strip() if company_elem else "Unknown Sender"
            job_info["description"] = soup.get_text(strip=True)[:1000]
            job_info["url"] = extract_url(content) or source_url
            job_info["source"] = "Email Scraping"
        
        logger.info(f"Parsed job listing: {job_info['title']} from {job_info['company']}")
        return job_info if any(keyword in job_info["title"].lower() or keyword in job_info["description"].lower() for keyword in JOB_KEYWORDS) else None
    except Exception as e:
        logger.error(f"Error parsing job listing from {source_type}: {str(e)}", exc_info=True)
        return None

def extract_url(text):
    """Extract the first URL from text content."""
    try:
        url_pattern = re.compile(r"https?://(?:[-\w.]|(?:%[\da-fA-F]{2}))+[-\w./?%&=]*")
        urls = url_pattern.findall(text)
        return urls[0] if urls else ""
    except Exception as e:
        logger.error(f"Error extracting URL: {str(e)}")
        return ""

async def save_job_to_vacante(job_info, bu):
    """Save parsed job information to Vacante model with async operation."""
    try:
        vacante = Vacante(
            nombre=job_info["title"],
            empresa=job_info["company"],
            descripcion=job_info["description"],
            url=job_info["url"],
            business_unit=bu,
            fuente=f"{job_info['source']} Parsing",
            fecha_publicacion=timezone.now(),
        )
        await sync_to_async(vacante.save)()
        logger.info(f"Saved parsed vacante: {job_info['title']}")
        return vacante
    except Exception as e:
        logger.error(f"Error saving parsed vacante: {str(e)}", exc_info=True)
        return None

class Parser:
    """Clase base para parsing de contenido HTML/JSON."""
    
    def __init__(self):
        self._cache = {}
        self._cache_ttl = 3600  # 1 hora
        self.nlp = NLPProcessor()
        
    def _get_cached_result(self, key: str) -> Optional[Dict]:
        """Obtiene resultado del caché si existe y es válido."""
        if key in self._cache:
            cached = self._cache[key]
            if time.time() - cached['timestamp'] < self._cache_ttl:
                return cached['data']
        return None
        
    def _set_cached_result(self, key: str, data: Dict):
        """Guarda resultado en caché."""
        self._cache[key] = {
            'data': data,
            'timestamp': time.time()
        }
        
    def parse_job_listing(self, content: str, url: str, source_type: str = "web") -> Dict:
        """
        Parsea una oferta de trabajo y extrae información relevante.
        
        Args:
            content: Contenido HTML/JSON de la oferta
            url: URL de la oferta
            source_type: Tipo de fuente (web, api, etc)
            
        Returns:
            Dict con la información extraída
        """
        cache_key = f"parse_{url}"
        cached = self._get_cached_result(cache_key)
        if cached:
            return cached
            
        try:
            if source_type == "json":
                data = json.loads(content)
                result = self._parse_json_job(data)
            else:
                result = self._parse_html_job(content)
                
            # Enriquecer con información adicional
            result.update({
                'url': url,
                'source_type': source_type,
                'parsed_at': datetime.now().isoformat()
            })
            
            # Guardar en caché
            self._set_cached_result(cache_key, result)
            
            return result
            
        except Exception as e:
            logger.error(f"Error parsing job listing: {e}")
            return {
                'url': url,
                'error': str(e)
            }
            
    def _parse_html_job(self, content: str) -> Dict:
        """Parsea una oferta de trabajo en formato HTML."""
        # Usar trafilatura para extraer el contenido principal
        main_content = trafilatura.extract(content) or content
        soup = BeautifulSoup(main_content, "html.parser")
        
        # Extraer información básica
        title = self._extract_title(soup)
        description = self._extract_description(soup)
        location = self._extract_location(soup)
        company = self._extract_company(soup)
        skills = self._extract_skills(soup)
        salary = self._extract_salary(soup)
        posted_date = self._extract_posted_date(soup)
        
        return {
            'title': title,
            'description': description,
            'location': location,
            'company': company,
            'skills': skills,
            'salary': salary,
            'posted_date': posted_date
        }
        
    def _parse_json_job(self, data: Dict) -> Dict:
        """Parsea una oferta de trabajo en formato JSON."""
        return {
            'title': data.get('title', ''),
            'description': data.get('description', ''),
            'location': data.get('location', {}).get('name', ''),
            'company': data.get('company', {}).get('name', ''),
            'skills': data.get('skills', []),
            'salary': data.get('salary', {}),
            'posted_date': data.get('posted_date', '')
        }
        
    def _extract_title(self, soup: BeautifulSoup) -> str:
        """Extrae el título de la oferta."""
        title_elem = (
            soup.find('h1') or 
            soup.find('h2') or 
            soup.find(class_=re.compile(r'title|heading', re.I))
        )
        return title_elem.get_text(strip=True) if title_elem else ''
        
    def _extract_description(self, soup: BeautifulSoup) -> str:
        """Extrae la descripción de la oferta."""
        desc_elem = (
            soup.find(class_=re.compile(r'description|content|details', re.I)) or
            soup.find('div', {'role': 'main'}) or
            soup.find('main')
        )
        return desc_elem.get_text(strip=True) if desc_elem else ''
        
    def _extract_location(self, soup: BeautifulSoup) -> str:
        """Extrae la ubicación de la oferta."""
        loc_elem = (
            soup.find(class_=re.compile(r'location|city|place', re.I)) or
            soup.find('span', {'itemprop': 'jobLocation'})
        )
        return loc_elem.get_text(strip=True) if loc_elem else ''
        
    def _extract_company(self, soup: BeautifulSoup) -> str:
        """Extrae el nombre de la empresa."""
        comp_elem = (
            soup.find(class_=re.compile(r'company|employer', re.I)) or
            soup.find('span', {'itemprop': 'hiringOrganization'})
        )
        return comp_elem.get_text(strip=True) if comp_elem else ''
        
    def _extract_skills(self, soup: BeautifulSoup) -> List[str]:
        """Extrae las habilidades requeridas."""
        skills = []
        
        # Buscar en secciones comunes de habilidades
        skill_sections = soup.find_all(class_=re.compile(r'skills|requirements|qualifications', re.I))
        for section in skill_sections:
            items = section.find_all(['li', 'p'])
            skills.extend([item.get_text(strip=True) for item in items])
            
        # Buscar en la descripción
        if not skills:
            desc = self._extract_description(soup)
            skills = self._extract_skills_from_text(desc)
            
        return list(set(skills))  # Eliminar duplicados
        
    def _extract_skills_from_text(self, text: str) -> List[str]:
        """Extrae habilidades de un texto usando expresiones regulares."""
        skills = []
        
        # Patrones comunes de habilidades
        patterns = [
            r'(?i)experience with\s+([^.,]+)',
            r'(?i)knowledge of\s+([^.,]+)',
            r'(?i)proficient in\s+([^.,]+)',
            r'(?i)expertise in\s+([^.,]+)',
            r'(?i)strong\s+([^.,]+)\s+skills',
        ]
        
        for pattern in patterns:
            matches = re.finditer(pattern, text)
            skills.extend([m.group(1).strip() for m in matches])
            
        return list(set(skills))
        
    def _extract_salary(self, soup: BeautifulSoup) -> Dict:
        """Extrae información del salario."""
        salary_elem = (
            soup.find(class_=re.compile(r'salary|compensation', re.I)) or
            soup.find('span', {'itemprop': 'baseSalary'})
        )
        
        if not salary_elem:
            return {}
            
        salary_text = salary_elem.get_text(strip=True)
        
        # Extraer rango de salario
        range_match = re.search(r'(\d+(?:,\d+)*)\s*-\s*(\d+(?:,\d+)*)', salary_text)
        if range_match:
            return {
                'min': int(range_match.group(1).replace(',', '')),
                'max': int(range_match.group(2).replace(',', '')),
                'currency': 'EUR' if '€' in salary_text else 'USD'
            }
            
        # Extraer salario fijo
        fixed_match = re.search(r'(\d+(?:,\d+)*)', salary_text)
        if fixed_match:
            return {
                'amount': int(fixed_match.group(1).replace(',', '')),
                'currency': 'EUR' if '€' in salary_text else 'USD'
            }
            
        return {}
        
    def _extract_posted_date(self, soup: BeautifulSoup) -> str:
        """Extrae la fecha de publicación."""
        date_elem = (
            soup.find(class_=re.compile(r'date|posted|published', re.I)) or
            soup.find('time') or
            soup.find('span', {'itemprop': 'datePosted'})
        )
        return date_elem.get_text(strip=True) if date_elem else ''

    def parse_education(self, text: str) -> Dict[str, Any]:
        """
        Parsea la sección de educación del CV.
        """
        education_data = {
            'universities': [],
            'degrees': [],
            'years': []
        }
        
        # Extraer universidades
        universities = self.nlp.extract_universities(text)
        education_data['universities'] = universities
        
        # Extraer grados
        degrees = self._extract_degrees(text)
        education_data['degrees'] = degrees
        
        # Extraer años
        years = self._extract_years(text)
        education_data['years'] = years
        
        return education_data
    
    def _extract_degrees(self, text: str) -> List[Dict[str, Any]]:
        """
        Extrae grados académicos.
        """
        # TODO: Implementar extracción de grados
        return []
    
    def _extract_years(self, text: str) -> List[Dict[str, Any]]:
        """
        Extrae años de estudio.
        """
        # TODO: Implementar extracción de años
        return []
    